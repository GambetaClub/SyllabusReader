import os
import re
from ics_converter import ICSConverter
import dateutil.parser
import pandas as pd
from docx import Document

class Reader:
    def __init__(self, dir=None, syllabi=None, calendar=None):
        self.__dir = dir
        self.__syllabi = syllabi
        self.__calendar = calendar

    def set_dir(self, dir):
        self.__dir = dir

    def get_dir(self):
        return self.__dir

    def get_filenames(self):
        return [key for key in self.__tables.keys()]

    def set_syllabi(self, syllabi):
        self.__syllabi = syllabi

    def get_syllabi(self):
        return self.__syllabi

    def get_calendar(self):
        return self.__calendar

    def get_docx_course(self, document):
        """
        Takes a docx (docx.Document) and returns
        the name of course the syllabus belongs to
        """
        # for paragraph in document.paragraphs:
        #     print(paragraph.text)

    def check_date(self, d):
        """
        Takes a string date and returns
        the a stringuniform date format if
        possible. Otherwise, it returns None 
        """
        try:
            d = dateutil.parser.parse(d)
            return d
        except:
            return None

    def insert_to_start(self, assignment, course):
        return course + " - " + assignment

    def spec(self, s):
        """
        Takes a string. Returns False is the string is
        made by more than just special characters.
        Otherwise, it returns True.
        """
        if not re.match(r'^[_\W]+$', s) and s != "":
            return s
        else:
            return None

    def check_syllabi(self):
        """
        Set the value of syllabi to be 
        a dictionary of the valid of the 
        documents with a valid format.
        """
        syllabi = self.get_syllabi()
        new_syllabi = dict()
        for key, value in syllabi.items():
            if value is None:
                continue
            new_syllabi[key] = value
        self.set_syllabi(new_syllabi)

    def convert_dates(self, df):
        """
        Convert the dates of the syllabi's
        dataframes into a monotonous format
        and sets the new dictionary for the 
        object attribute.
        """
        df['Date'] = df['Date'].map(lambda d: self.check_date(d))
        df = df[df.Date.notnull()]
        return df 

    def convert_assignments(self, df, filename):
        """
        Converts the dataframe assignments 
        values to None if the value is either
        only spaces or symbol characters. 
        """
        with pd.option_context('mode.chained_assignment', None):
            df['Assignments'] = df['Assignments'].map(lambda s: self.spec(s))
            df = df[df.Assignments.notnull()]
            df['Assignments'] = df["Assignments"].map(lambda s: self.insert_to_start(s, filename))
            return df

    def recognize_fields(self, df):
        """
        Accepts a dataframe and returns 
        the dataframe only with the fields
        "Assignments", "Week", and "Date".
        """
        if isinstance(df, pd.DataFrame):
            if "Assignments" in list(df) or "Week" in list(df) or "Date" in list(df):
                df = df[["Assignments", "Week", "Date"]]
            else:
                return pd.DataFrame()
        return df
        
    def read_docx_table(self, document, n_headers=1):
        """
        Reads a document's tables (docx Document) 
        and returns with dataframes that represent
        the calendar in the syllabus. If the syllabus
        doesn't contain  any table with the format, it returns None. 
        """ 
        tables = document.tables
        if not tables:
            return None

        dfs = list()
        for table in tables:
            data = [[cell.text for cell in row.cells] for row in table.rows]
            df = pd.DataFrame(data)

            if n_headers == 1:
                df = df.rename(columns=df.iloc [0]).drop (df.index [0]).reset_index(drop=True)
            
            df = self.recognize_fields(df)
            if not df.empty:  
                dfs.append(df)
        if not dfs:
            return None
        return dfs[0]

    def load_syllabi(self, dir=None):
        """
        Accepts a dir and reads their tables
        with the calendar information for each
        single docx file in the directory.
        At the end it sets the dataframes as the
        object syllabi.
        """
        if dir != None:
            self.set_dir(dir)
        syllabi = dict()
        dir_name = self.get_dir()
        for filename in os.listdir(dir_name):
            if filename.startswith("~$"):
                continue
            if filename.endswith(".docx"):
                document = Document(os.path.join(dir_name, filename))
                syllabi[os.path.splitext(filename)[0]] = self.read_docx_table(document)
        self.set_syllabi(syllabi)
        self.check_syllabi()

        syllabi = self.get_syllabi()
        for syllabus in syllabi:
            df = syllabi[syllabus]
            df = self.convert_dates(df)
            df = self.convert_assignments(df, syllabus)
            syllabi[syllabus] = df
        self.set_syllabi(syllabi)

    def get_files_dir(self, args):
        """
        Interface helper function that resolves the problem 
        of getting the directory name in which they 
        files (docx or csv) are in.
        """
        files_dir = os.getcwd()
        # If the directory was passed as an argument then it resolves
        if len(args) == 2:
            files_dir = os.path.join(files_dir, args[1])
        # If not, then it asks for the user to type the name
        else:
            print("What's the dir's name in which the csv files are?")
            print("Press enter if it's in the current dir.")
            dir_name = input()
            files_dir = os.path.join(files_dir, dir_name)
        return files_dir

    def convert_docx_to_csv(self, files_dir):
        # Setting the docx directory
        self.set_dir(files_dir)
        
        # Transforming docx tables to dataframes
        self.load_syllabi()
        
        # Getting the dataframes stored in the instance
        syllabi = self.get_syllabi()
        for syllabus in syllabi:
            # Creating a csv file based on the instance syllabi's dataframe
            syllabi[syllabus].to_csv(f"{syllabus}.csv", encoding='utf-8', index=False)
            
            # Moving files to the same directory with other files
            old_file_abs_path = os.path.join(os.getcwd(), f"{syllabus}.csv")
            new_file_abs_path = os.path.join(os.getcwd(), files_dir, f"{syllabus}.csv")
            os.rename(old_file_abs_path, new_file_abs_path)

    def convert_csv_to_ics(self, files_dir):
        # Creating ics converter object and passing the csv files dir
        converter = ICSConverter(files_dir)

        # Creating a list of only the name files that end with csv
        csv_filenames = [filename for filename in os.listdir(files_dir) if filename.endswith('.csv')]

        # Defining the dir absolute path from where to get the csv files
        abs_dir_path = os.path.abspath(files_dir)

        if csv_filenames:
            for filename in csv_filenames:
                # Pass csv file to the converter
                converter.readCSV(os.path.join(abs_dir_path, filename))
            
            # Convert instance saved csv files to ics format
            converter.exportICS()
        else:
            print("There are no csv files in the listed dir.")
            return False

    def display_interface(self, args):
        print("Options:")
        print(f"1: Convert docx to csv format.")
        print(f"2: Convert csv to ics format.")
        print(f"3: Convert docx to ics format.")
        option = input()
        files_dir = self.get_files_dir(args)
        
        if option == '1':
            self.convert_docx_to_csv(files_dir)
        elif option == '2':
            self.convert_csv_to_ics(files_dir)
        elif option == '3':
            self.convert_docx_to_csv(files_dir)
            self.convert_csv_to_ics(files_dir)